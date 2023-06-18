from django.http import HttpResponse
from django.http.response import JsonResponse
from django.shortcuts import render
from rest_framework import viewsets
from basket.basket import Basket
from orders.serializer import OrderDeliveryOptionSerializer, OrderItemsSerializer, OrderSerializer
from store.models import Product
from django.db.models import Count, Sum
from django.http import FileResponse
from docx import Document
import io
from django.contrib import messages
from datetime import datetime, timedelta

from .models import Order, OrderDeliveryOption, OrderItem


def add(request):
    basket = Basket(request)
    if request.POST.get('action') == 'post':

        order_key = request.POST.get('order_key')
        user_id = request.user.id
        baskettotal = basket.get_total_price()

        # Check if order exists
        if Order.objects.filter(order_key=order_key).exists():
            pass
        else:
            order = Order.objects.create(user_id=user_id, full_name='name', address1='add1',
                                address2='add2', total_paid=baskettotal, order_key=order_key)
            order_id = order.pk

            for item in basket:
                print(item)
                OrderItem.objects.create(order_id=order_id, product=item['product'], price=item['price'], quantity=item['qty'])

        response = JsonResponse({'success': 'Return something'})
        return response


def payment_confirmation(data):
    Order.objects.filter(order_key=data).update(billing_status=True)


def user_orders(request):
    orders = Order.objects.all()

    print(orders)
    return orders
    

def admin_orders(request):
    orders = Order.objects.all()
    print(request.method)
    if request.method == "POST":
        deta =  request.POST['search_data']

        print("---------->",deta)

        if(deta.isnumeric()):

            if deta == "":
                print("nothing")
                messages.error(request, 'Please insert something')
            else:
                orders = Order.objects.filter(id = deta)
            return render(request, "admin/orders.html", {"orders": orders})
        else:
            messages.error(request, 'Wrong input')

    print(orders)
    return render(request, "admin/orders.html", {"orders": orders})
        
def admin_orders_report(request):
    print(request.method)
    if request.method == "POST":

        return render(request, "admin/admin_orders_report.html")
    
    return render(request, "admin/admin_orders_report.html")

def download_txt(request):
    # Create some text content
    content = "This is some example text.\nIt can have multiple lines.\n"

    orders = Order.objects.all()

    for order in orders:
        print(order.id)

        data = '{orderId: '+ str(order.id)+ ',full name: '+order.full_name+',email: '+order.email+',order date:'+ str(order.created)+',address: '+order.address1+ order.address2+ order.postal_code+',total paid:'+ str(order.total_paid)+'},\n'
        content += data
        print("====================>afrer",content)

    messege = content

    # Create the HttpResponse object with the content and file type header
    response = HttpResponse(messege, content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename="example.txt"'

    return response

class OrdersListView(viewsets.ModelViewSet):
    serializer_class = OrderSerializer
    queryset = Order.objects.all()


class OrderItemsListView(viewsets.ModelViewSet):
    serializer_class = OrderItemsSerializer
    queryset = OrderItem.objects.all()

class OrderDeliveryOptionList(viewsets.ModelViewSet):
    serializer_class = OrderDeliveryOptionSerializer
    queryset = OrderDeliveryOption.objects.all()


def overall_summary(request):

    if request.method == "POST":

        data = Product.objects.values('shop_retail').annotate(count=Count('id'))
        most_ordered_products = OrderItem.objects.values('product_id').annotate(order_count=Sum('quantity')).order_by('-order_count')[:10]

        print("================>",most_ordered_products)
        list_of_ids = []
        for mop in most_ordered_products:
            print(mop['product_id'])
            list_of_ids.append(mop['product_id'])

        products = Product.objects.filter(id__in=list_of_ids)

        content = ''

        for prod in products:
            print(prod.id)

            data = '{product_Id: '+ str(prod.id)+ ',product name: '+prod.product_name+',price: '+prod.price+',product image:'+ prod.product_image+',Retails: '+prod.shop_retail+'},\n______________________________________________________________\n'
            content += data
            print("====================>afrer",content)

        messege = content
        
        # create a new Word document using the python-docx library
        document = Document()
        document.add_heading('Top Ten most ordered products', 0)
        document.add_paragraph(messege)

        # render the document to a BytesIO object
        document_buffer = io.BytesIO()
        document.save(document_buffer)

        # create a HttpResponse object with the document content and the appropriate headers
        response = HttpResponse(document_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Top-ten-products.docx'

        return response

    data = Product.objects.values('shop_retail').annotate(count=Count('id'))
    most_ordered_products = OrderItem.objects.values('product_id').annotate(order_count=Sum('quantity')).order_by('-order_count')[:10]

    print("================>",most_ordered_products)
    list_of_ids = []
    for mop in most_ordered_products:
        print(mop['product_id'])
        list_of_ids.append(mop['product_id'])

    products = Product.objects.filter(id__in=list_of_ids)
    print(products)
    context = {
        'products':products,
        'data': data,
    }
    
    return render(request, 'admin/overall_summary.html', context)

def filter_admin_orders_by_date(request):

    if request.method == "POST":

        deta1 =  request.POST['from']
        deta2 =  request.POST['to']
        print('======================>my dates',deta1,deta2)
        try:
            date_obj_from = datetime.strptime(deta1, '%Y-%m-%d')
            date_obj_to = datetime.strptime(deta2, '%Y-%m-%d')
            # Do something with the valid date object
            if date_obj_to > date_obj_from:
                print('yes')
                orders = Order.objects.filter(billing_status=True).filter(created__range=[date_obj_from,date_obj_to])
                return render(request, "admin/admin_orders_report.html", {"orders": orders})
            elif date_obj_to == date_obj_from:
                messages.error(request, 'Please use the minimum of one day e.g: 2022/01/01 - 2022/01/02')
                return render(request, "admin/admin_orders_report.html")
            else:
                print('incorrect dates')
                messages.error(request, 'Incorrect date, From date must be less than To date')
                return render(request, "admin/admin_orders_report.html")
        except ValueError:
            # Handle the case where the input is not a valid date
            print('Invalid date')
            messages.error(request, 'Invalid date')
            return render(request, "admin/admin_orders_report.html")
    
    return render(request, "admin/admin_orders_report.html")